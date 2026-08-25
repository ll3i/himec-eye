# -*- coding: utf-8 -*-
"""
공모전 제출서류(붙임4 출품작) 자동 작성

원본 양식 docx를 복사한 뒤, 붙임4 영역의 설명 셀에 제안 내용을 채워 넣는다.
붙임1~3(참가신청서/서약서/개인정보 동의서)은 개인정보가 필요하므로 비워 둔다.
"""
import os, sys, os, json, shutil, argparse
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
ROOT = Path(__file__).resolve().parents[1]
FORM = ROOT.parent / "제1회 HIMEC AI 활용 아이디어 공모전" / "제1회 HIMEC AI 활용 아이디어 공모전" \
       / "2. 붙임+제출서류(제1회 HIMEC AI 활용 아이디어 공모전).docx"

TITLE = 'HIMEC EYE — 직무별 비전 AI 관제 : 현장 사진 1장에서 안전·품질 지적사항과 근거 조항·조치까지'

SUMMARY = "미화·보안·생산도급·물류·주차처럼 사람이 직접 수행하는 현장 서비스는 점검이 순회 인력의 눈에 의존한다. 같은 사진도 보는 사람에 따라 판정이 갈리고 기록은 사후 수기로 남는다. 직무마다 보는 대상과 판정 기준이 다르므로 하나의 만능 모델 대신 직무별 전용 모델을 두고, 공통 룰엔진이 근거 조항·위험도·조치사항을 붙인다. AI는 '보는 일'만 하고 법규 판정은 결정론적 규칙이 맡는다. 6개 직무를 실제로 학습·판독하는 프로토타입까지 구현했다."

BACKGROUND = """■ 현황

제안자는 삼구아이앤씨에서 AI 엔지니어로 일하고 있습니다. 삼구아이앤씨는 미화·보안·생산도급·물류·주차 운영처럼 사람이 직접 수행하는 현장 서비스를 여러 사업장에서 동시에 운영합니다.

한 사업장 안에서도 생산 라인, 자재 야드, 기계실, 주차 구획이 각각 다른 위험과 다른 품질 기준을 갖습니다. 그리고 이 모든 곳의 상태 확인은 결국 순회하는 사람의 눈과, 그 사람이 찍어 온 사진에 남습니다.

■ 문제점

① 판독이 사람에 묶여 있습니다. 안전 점검을 나간 사람은 보호구를 보고, 설비 점검을 나간 사람은 배관을 봅니다. 같은 사진에 둘 다 찍혀 있어도 자기 목적에 해당하는 것만 보고 나머지는 지나갑니다.

② 기준이 속인화됩니다. 동일한 부식·균열 사진을 두고도 담당자의 경력에 따라 경미/보통/중대 판정이 갈립니다. 신입은 어떤 조항을 근거로 지적해야 하는지 매번 찾아야 하고, 그 근거가 문서에 남지 않으면 나중에 다툼의 소지가 됩니다.

③ 기록이 사후에 수기로 만들어집니다. 현장에서 찍고, 사무실에 돌아와 사진을 정리하고, 지적사항을 글로 옮기고, 근거 조항을 찾아 붙입니다. 사진 수가 늘수록 이 뒷단이 병목이 됩니다.

④ 직무마다 판정 기준이 다릅니다. 안전은 놓치면 사고로 이어지므로 의심 수준에서도 올려야 하고, 품질 불량은 오탐이 잦으면 라인이 서므로 확실할 때만 올려야 합니다.

■ 개선 필요 사항

이 문제는 삼구아이앤씨만의 것이 아닙니다. HIMEC이 60년간 축적한 통합 엔지니어링 역량으로 수행하는 시공·품질관리 현장도 같은 구조입니다 — 현장 사진은 쌓이는데, 판독과 문서화는 여전히 사람의 눈과 손에 남아 있습니다. 본 공모전 ② 시공·품질관리 분야가 든 예시 자체가 '현장사진 AI분석으로 하자·안전 판독'입니다."""

PROPOSAL = """■ 제안 내용

현장 사진 1장을 입력하면, 그 사진이 속한 직무의 전용 모델이 판독하고 공통 룰엔진이 근거 조항·위험도·조치사항을 붙여 지적사항을 산출합니다.

핵심은 '하나의 만능 모델'을 만들지 않았다는 점입니다. 2차전지 셀의 극성 반전과 지게차 협착 위험은 찍는 대상도, 판정 기준도, 근거 법규도 다릅니다. 직무마다 전용 모델과 전용 규칙을 가진 에이전트를 두고, 오케스트레이터가 결과를 위험도 순으로 합칩니다.

[검증한 6개 직무] 안전관리(보호구·연기) / 시설·설비 관리(부식·균열·케이블·용접) / 생산 품질검사(냉납·조립 불량·단락·표면) / 2차전지 셀 선별(극성) / 물류·운송 관리(지게차·작업자·섀시) / 주차 운영관리(차종별 점유)

■ 적용할 AI 기술 및 활용 방안

1단계 지각 — 분류가 아니라 객체탐지를 씁니다. '어디에 무엇이 있다'가 나와야 지적 위치를 사진대지에 표시하고, 객체 사이의 관계를 따질 수 있습니다.

2단계 판정 — 법규 판정에 LLM을 쓰지 않았습니다. 조항을 잘못 인용한 검측 문서는 문서로서의 가치를 잃기 때문입니다. 코드로 고정된 규칙 21개가 근거 조항·위험도 4단계·조치사항을 함께 갖습니다.
 · 항목별 임계값 — 안전모 미착용 0.25(재현율 우선), 연기 0.55(정밀도 우선). 오경보를 남발하면 현장이 경보를 무시합니다. 역극 셀 0.25 / 정상 셀 0.40 처럼 방향이 반대인 정책도 씁니다.
 · 관계 추론 — 지게차 박스를 좌우 35%·상하 17.5% 확장한 영역에 작업자 박스가 35% 이상 겹치면 협착 위험. 각각 검출하는 것만으로는 나오지 않는 판정입니다. 보호구도 '작업자는 있는데 머리 구간에 안전모가 없다'는 관계로 추론합니다.
 · 정량 등급 — 검출 면적비로 경미(1% 미만)/보통(1~5%)/중대(5% 이상).

3단계 산출 — 지적사항(규정번호·위험도·근거 조항·조치사항·좌표), 보호구 준수율, 이전 점검 대비 해소·잔존·신규 판정.

■ 실증 성능 (자체 test split · GPU 없이 CPU 학습)
{PERF}

■ 기대효과

직무 경계를 넘어 한 장에서 동시에 판독하므로 누락이 줄고, 판정이 코드로 고정되어 담당자 경력에 따라 결과가 달라지지 않습니다. 근거 조항이 자동 인용되어 지적의 정당성이 문서에 남습니다. 새 직무는 해당 모델과 규칙만 붙이면 되어 확산이 쉽습니다.

프로토타입은 Ultralytics YOLO11(AGPL-3.0)로 학습했습니다. 사내 상용 적용 시 상용 라이선스 취득 또는 Apache-2.0 계열 대체 모델 재학습이 필요합니다. 학습 데이터는 라이선스가 명확한 공개 데이터셋(CC/CC0)만 사용했습니다."""

ATTACH = """[온라인 시연 - 심사위원이 직접 조작 가능]
  https://himec-eye-panel.vercel.app        ← 주 시연 페이지 (화면 5종)
  · 관제 화면        현장 화면 위에 판정 결과를 그대로 표시합니다. 감시 → 스캔 → 검출 → 경보 확정 4단계가 순환하며, 위험도에 따라 긴급(적색)과 부적합(주황)으로 구분됩니다. 지게차 협착 위험처럼 '두 객체의 관계'로 나온 판정은 관련 객체를 함께 표시합니다.
  · 현장 재현 영상   현장 카메라 화면과 판독 화면을 나란히 두고 판독 과정을 재현합니다.
  · 판독 결과        6개 직무 · 시험 사진 23장. 검출 항목별 신뢰도를 임계값과 함께 표시합니다.
  · 판정 작동        직무별 검출 클래스와 임계값 설정 근거, 판정 규칙 21개(근거 조항·위험도 포함). '미채택 표시'를 켜면 임계값 미만으로 걸러진 검출까지 확인할 수 있습니다.
  · 분석 리포트      시험 사진 23장의 판독 결과를 정리한 인쇄용 리포트.
  판정 로직(항목별 임계값 정책 · 관계 추론 · 위험도 산정 · 근거법규 매핑)은 파이썬 룰엔진(rules.py)과 같은 조건으로 브라우저에서 실행됩니다.

[시연 영상]
  https://youtu.be/__________          ← 관제 화면 36초 (4개 현장 · 판독부터 경보 확정까지)

[소스코드]
  https://github.com/ll3i/himec-eye     ← 데이터 수집 · 학습 · 판독 엔진 · 룰엔진 · 시연 페이지 빌더

[제출 파일]
· 프로토타입 소스코드 : 데이터 수집기 / 학습 스크립트 / 판독 엔진 / 판정 룰엔진 / 문서 자동생성기 / 시연 페이지 빌더
· 실행 결과 : 현장사진 판독 결과 시각화 이미지, 자동 생성된 점검 문서(HTML), 판독 과정 시연 영상(GIF)
· 주 시연 페이지 사본 : 오프라인에서도 열 수 있는 정적 파일 일체
· 데이터셋 명세 : 수집 데이터의 구성 및 출처·라이선스 명세
· 성능 리포트 : 직무별 모델 학습 결과(mAP / Precision / Recall) 및 클래스별 성능표
· README : 프로젝트 구조 및 재현 실행 방법"""


# 출품 대상 6개 직무. 본문이 "검증한 6개 직무"라고 쓰므로 목록이 어긋나면 안 된다.
SUBMITTED = [("safety", "안전관리"), ("defect", "시설·설비 관리"),
             ("product", "생산 품질검사"), ("polarity", "2차전지 셀 선별"),
             ("logistics", "물류·운송 관리"), ("parking", "주차 운영관리")]


def build_perf(perf_json):
    """붙임4 설명칸에 들어갈 실증 성능 블록.

    설명칸은 표 안의 한 셀이라 쪽을 넘겨 흐르지 못한다. 줄 수를 아끼려고
    직무별 한 줄 대신 세 개씩 묶어 두 줄로 만든다.
    """
    perf_json = Path(perf_json)
    if not perf_json.exists():
        return ""
    reps = {r["task"]: r for r in json.loads(perf_json.read_text(encoding="utf-8"))}
    cells = [f"{ko} {reps[t]['overall']['map50']*100:.1f}%"
             for t, ko in SUBMITTED if t in reps]
    if not cells:
        return ""
    rows = [" · " + "   · ".join(cells[i:i + 3]) for i in range(0, len(cells), 3)]
    eps = sorted({reps[t].get("epochs") for t, _ in SUBMITTED
                  if t in reps and reps[t].get("epochs")})
    span = f"{eps[0]}~{eps[-1]}" if len(eps) > 1 else (str(eps[0]) if eps else "-")
    rows.append(f"   (mAP@50 · 자체 test split · GPU 없이 CPU {span} epoch 학습)")
    extra = len([k for k in reps if k not in dict(SUBMITTED)])
    if extra:
        rows.append(f" · 위 6개 외에 {extra}개 직무(위생·미화·반도체·계기)도 같은 구조로 "
                    "학습해 두었으며, 데이터가 확보되는 대로 붙일 수 있습니다.")
    return "\n".join(rows)


def set_cell(cell, text, size=10):
    from docx.shared import Pt
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    first = True
    for line in text.split("\n"):
        if first:
            para = p
            first = False
        else:
            para = cell.add_paragraph()
        run = para.add_run(line)
        run.font.size = Pt(size)
        run.font.name = "맑은 고딕"
    return cell


APPLICANT_EMAIL = os.environ.get("APPLICANT_EMAIL", "")
FIELD = "시공·품질관리"


def fill_application(doc):
    """붙임1 참가신청서 중 개인정보가 아닌 칸만 채운다.

    신청자명·연락처·소속·팀원·서명·날짜와 붙임2·붙임3은 본인이 직접 써야 하므로
    건드리지 않는다.
    """
    done = []
    for tb in doc.tables:
        for row in tb.rows:
            cells, last = [], None
            for c in row.cells:
                if c._tc is not last:
                    cells.append(c)
                    last = c._tc
            if not cells:
                continue
            head = cells[0].text.strip()
            body = cells[-1].text.strip()

            if "설계·엔지니어링" in body and ("□ " + FIELD) in body:
                set_cell(cells[-1], body.replace("□ " + FIELD, "■ " + FIELD))
                done.append("공모분야 체크")
            elif head == "연 락 처":
                for i, c in enumerate(cells):
                    if c.text.strip() == "E-Mail" and i + 1 < len(cells) \
                            and not cells[i + 1].text.strip():
                        set_cell(cells[i + 1], APPLICANT_EMAIL)
                        done.append("E-Mail")
            elif head == "출품작 명칭" and not body:
                set_cell(cells[-1], TITLE)
                done.append("출품작 명칭")
            elif head.startswith("공모 내용"):
                set_cell(cells[-1], SUMMARY)
                done.append("공모 내용/요약")
    return done


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default=str(ROOT / "reports" / "HIMEC AI 활용 아이디어 공모전_출품작.docx"))
    ap.add_argument("--perf", default=str(ROOT / "reports" / "performance.md"))
    a = ap.parse_args()

    perf_txt = build_perf(Path(a.perf).with_suffix(".json"))

    from docx import Document
    Path(a.out).parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(FORM, a.out)
    doc = Document(a.out)

    tables = doc.tables
    print(f"tables in form: {len(tables)}")
    filled = 0
    for t in tables:
        try:
            head = t.rows[0].cells[0].text.strip()
        except Exception:
            continue
        body = t.rows[0].cells[-1].text
        if "현황, 문제점 및 개선 필요 사항" in body:
            set_cell(t.rows[0].cells[-1], BACKGROUND)
            filled += 1
        elif "적용할 AI 기술 및 활용 방안 설명" in body:
            set_cell(t.rows[0].cells[-1], PROPOSAL.replace("{PERF}", perf_txt))
            filled += 1

    # 붙임4 본문 문단 채우기
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        s = p.text.strip()
        if s.startswith("1. 출품작 명칭"):
            for r in list(p.runs)[1:]:
                r._element.getparent().remove(r._element)
            if p.runs:
                p.runs[0].text = "1. 출품작 명칭 : " + TITLE
            else:
                p.add_run("1. 출품작 명칭 : " + TITLE)
            filled += 1
        elif s.startswith("(3) 첨부 자료"):
            for q in paras[i + 1: i + 4]:
                if q.text.strip().startswith(":"):
                    for r in list(q.runs)[1:]:
                        r._element.getparent().remove(r._element)
                    body = ": PPT, 프로토타입, GitHub주소, 기타 등"
                    if q.runs:
                        q.runs[0].text = body
                    else:
                        q.add_run(body)
                    for line in ATTACH.split("\n"):
                        np = q.insert_paragraph_before(line) if False else None
                    # 첨부 목록은 별도 문단으로 이어붙임
                    parent = q._element.getparent()
                    ref = q._element
                    from docx.text.paragraph import Paragraph
                    import copy
                    for line in ATTACH.split("\n"):
                        newp = copy.deepcopy(q._element)
                        parent.insert(list(parent).index(ref) + 1, newp)
                        np = Paragraph(newp, q._parent)
                        for r in list(np.runs)[1:]:
                            r._element.getparent().remove(r._element)
                        if np.runs:
                            np.runs[0].text = line
                        ref = newp
                    filled += 1
                    break

    app = fill_application(doc)
    print("붙임1 자동 기입:", ", ".join(app) if app else "없음")

    doc.save(a.out)
    print(f"filled sections: {filled}")
    print("saved:", a.out)
    print("\n[남은 작업] 붙임1 참가신청서 / 붙임2 참가서약서 / 붙임3 개인정보 동의서 는 "
          "개인정보가 필요하므로 직접 작성해야 합니다.")
    print("[요약란] 붙임1 '공모 내용/요약(300자 이내)'에 넣을 문구:")
    print("-" * 60)
    print(SUMMARY)
    print("-" * 60, f"({len(SUMMARY)}자)")


if __name__ == "__main__":
    main()
